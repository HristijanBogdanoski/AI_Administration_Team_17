from copy import deepcopy
import re
from io import BytesIO

from sqlalchemy.orm import Session

from app.models.service import Service
from app.models.user import User
from app.models.service_document_template import ServiceDocumentTemplate

# PDF/DOCX generation
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import simpleSplit
from docx import Document
import os
from app.schemas.service_document_template import (
    ServiceDocumentTemplateCreate,
    ServiceDocumentTemplateUpdate,
)


def build_default_template_body(service_name: str) -> dict:
    return {
        "document_type": "service_application",
        "service_name": service_name,
        "sections": [
            {
                "title": "Податоци за апликантот",
                "fields": [
                    {"key": "full_name", "label": "Име и презиме", "value": ""},
                    {"key": "embg", "label": "ЕМБГ", "value": ""},
                    {"key": "address", "label": "Адреса", "value": ""},
                    {"key": "phone_number", "label": "Телефон", "value": ""},
                    {"key": "gender", "label": "Пол", "value": ""},
                    {"key": "email", "label": "Е-пошта", "value": ""},
                ],
            },
            {
                "title": "Детали за услугата",
                "fields": [
                    {"key": "service_name", "label": "Услуга", "value": service_name},
                    {"key": "notes", "label": "Забелешки", "value": ""},
                ],
            },
        ],
    }


def get_user_field_values(user: User) -> dict[str, str | None]:
    return {
        "full_name": user.full_name,
        "email": user.email,
        "embg": user.embg,
        "address": user.address,
        "phone_number": user.phone_number,
        "gender": user.gender,
    }


def apply_user_values_to_template(
    template_body: dict,
    user: User,
    selected_fields: list[str] | None = None,
) -> dict:
    filled_body = deepcopy(template_body)
    user_values = get_user_field_values(user)
    allowed_fields = set(selected_fields or user_values.keys())

    for section in filled_body.get("sections", []):
        for field in section.get("fields", []):
            key = field.get("key")
            if key in allowed_fields and key in user_values and user_values[key]:
                field["value"] = user_values[key]
            if key == "phone" and "phone_number" in allowed_fields and user_values.get("phone_number"):
                field["value"] = user_values["phone_number"]

    return filled_body


def render_template_document(template_body: dict) -> str:
    lines: list[str] = []
    service_name = template_body.get("service_name")
    if service_name:
        lines.append(str(service_name))
        lines.append("=" * len(str(service_name)))
        lines.append("")

    for section in template_body.get("sections", []):
        title = section.get("title")
        if title:
            lines.append(str(title))
        for field in section.get("fields", []):
            label = field.get("label", field.get("key", "Field"))
            value = field.get("value")
            display_value = value if value not in (None, "") else "________________"
            lines.append(f"{label}: {display_value}")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


def _register_font() -> str:
    font_name = "DejaVuSans"
    bold_name = "DejaVuSans-Bold"
    candidate_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "C:\\Windows\\Fonts\\DejaVuSans.ttf",
        "C:\\Windows\\Fonts\\arial.ttf",
        "C:\\Windows\\Fonts\\ARIALUNI.TTF",
    ]
    bold_paths = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
        "C:\\Windows\\Fonts\\arialbd.ttf",
    ]
    for p in candidate_paths:
        try:
            if os.path.exists(p):
                pdfmetrics.registerFont(TTFont(font_name, p))
                for bp in bold_paths:
                    if os.path.exists(bp):
                        pdfmetrics.registerFont(TTFont(bold_name, bp))
                        return font_name
                return font_name
        except Exception:
            continue
    return "Helvetica"


def text_to_pdf_bytes(text: str) -> bytes:
    # Parse plain text back into a minimal template_body for structured rendering
    lines = text.splitlines()
    service_name = lines[0] if lines else "Документ"
    body = {"service_name": service_name, "sections": [{"title": "", "fields": []}]}
    current_section = body["sections"][0]
    for line in lines[2:]:
        if not line.strip():
            continue
        if ":" in line:
            label, _, value = line.partition(":")
            current_section["fields"].append({"label": label.strip(), "value": value.strip()})
        else:
            new_section = {"title": line.strip(), "fields": []}
            body["sections"].append(new_section)
            current_section = new_section
    return template_body_to_pdf_bytes(body)


def template_body_to_pdf_bytes(template_body: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm

    font_name = _register_font()
    bold_name = font_name + "-Bold" if font_name == "DejaVuSans" else font_name + "-Bold"

    DARK_BLUE = colors.HexColor("#1B3A6B")
    GOLD = colors.HexColor("#D4A017")
    LIGHT_BLUE = colors.HexColor("#EFF6FF")
    GRAY = colors.HexColor("#64748b")
    WHITE = colors.white

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=15*mm, bottomMargin=20*mm
    )

    styles = {
        "service_name": ParagraphStyle("sn", fontName=font_name, fontSize=18, textColor=WHITE, spaceAfter=2),
        "subtitle": ParagraphStyle("sub", fontName=font_name, fontSize=9, textColor=colors.HexColor("#93c5fd"), spaceAfter=0),
        "section_title": ParagraphStyle("st", fontName=font_name, fontSize=11, textColor=DARK_BLUE, spaceBefore=10, spaceAfter=6, fontWeight="bold"),
        "label": ParagraphStyle("lbl", fontName=font_name, fontSize=9, textColor=GRAY),
        "value": ParagraphStyle("val", fontName=font_name, fontSize=11, textColor=colors.HexColor("#1e293b")),
        "footer": ParagraphStyle("ft", fontName=font_name, fontSize=8, textColor=GRAY, alignment=1),
    }

    from datetime import datetime
    service_name = template_body.get("service_name", "Документ")
    today = datetime.now().strftime("%d.%m.%Y %H:%M:%S")

    story = []

    # Header block
    header_data = [[
        Paragraph(service_name, styles["service_name"]),
    ]]
    header_table = Table(header_data, colWidths=[170*mm])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), DARK_BLUE),
        ("TOPPADDING", (0, 0), (-1, -1), 14),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 16),
        ("RIGHTPADDING", (0, 0), (-1, -1), 16),
        ("ROUNDEDCORNERS", [6, 6, 6, 6]),
    ]))
    story.append(header_table)

    # Gold accent line
    story.append(HRFlowable(width="100%", thickness=3, color=GOLD, spaceAfter=12))

    # Sections
    for section in template_body.get("sections", []):
        title = section.get("title", "")
        fields = section.get("fields", [])
        if not fields:
            continue

        if title:
            section_header = Table([[Paragraph(f"  {title}", styles["section_title"])]], colWidths=[170*mm])
            section_header.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BLUE),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("ROUNDEDCORNERS", [4, 4, 4, 4]),
            ]))
            story.append(section_header)
            story.append(Spacer(1, 6))

        # Fields in 2-column grid
        field_rows = []
        for i in range(0, len(fields), 2):
            row = []
            for field in fields[i:i+2]:
                label = field.get("label", field.get("key", ""))
                value = field.get("value") or ""
                display = value if value.strip() else "________________"
                cell_content = [
                    Paragraph(label, styles["label"]),
                    Paragraph(display, styles["value"]),
                ]
                row.append(cell_content)
            if len(row) == 1:
                row.append("")
            field_rows.append(row)

        if field_rows:
            field_table = Table(field_rows, colWidths=[82*mm, 82*mm], hAlign="LEFT")
            field_table.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("LINEBELOW", (0, 0), (-1, -1), 0.5, colors.HexColor("#e2e8f0")),
            ]))
            story.append(field_table)
            story.append(Spacer(1, 8))

    # Footer
    story.append(Spacer(1, 10))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e2e8f0"), spaceAfter=8))
    story.append(Paragraph(f"Влада на Република Северна Македонија  •  е-Влада Портал  •  Генерирано: {today}", styles["footer"]))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def text_to_docx_bytes(text: str) -> bytes:
    lines = text.splitlines()
    service_name = lines[0] if lines else "Документ"
    body = {"service_name": service_name, "sections": [{"title": "", "fields": []}]}
    current_section = body["sections"][0]
    for line in lines[2:]:
        if not line.strip():
            continue
        if ":" in line:
            label, _, value = line.partition(":")
            current_section["fields"].append({"label": label.strip(), "value": value.strip()})
        else:
            new_section = {"title": line.strip(), "fields": []}
            body["sections"].append(new_section)
            current_section = new_section
    return template_body_to_docx_bytes(body)


def template_body_to_docx_bytes(template_body: dict) -> bytes:
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from datetime import date

    DARK_BLUE = RGBColor(0x1B, 0x3A, 0x6B)
    GOLD = RGBColor(0xD4, 0xA0, 0x17)
    LIGHT_BLUE = RGBColor(0xEF, 0xF6, 0xFF)
    GRAY = RGBColor(0x64, 0x74, 0x8B)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(2)

    service_name = template_body.get("service_name", "Документ")
    today = date.today().strftime("%d.%m.%Y")

    # Header paragraph with dark blue background
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(service_name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Set background shading
    pPr = header_para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "1B3A6B")
    pPr.append(shd)
    header_para.paragraph_format.space_before = Pt(8)
    header_para.paragraph_format.space_after = Pt(8)
    header_para.paragraph_format.left_indent = Cm(0.3)

    # Gold line (using a thick border)
    gold_para = doc.add_paragraph()
    gold_pPr = gold_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4A017")
    pBdr.append(bottom)
    gold_pPr.append(pBdr)
    gold_para.paragraph_format.space_after = Pt(10)

    # Sections
    for section in template_body.get("sections", []):
        title = section.get("title", "")
        fields = section.get("fields", [])
        if not fields:
            continue

        if title:
            sec_para = doc.add_paragraph()
            sec_run = sec_para.add_run(f"  {title}")
            sec_run.bold = True
            sec_run.font.size = Pt(11)
            sec_run.font.color.rgb = DARK_BLUE
            sec_pPr = sec_para._p.get_or_add_pPr()
            sec_shd = OxmlElement("w:shd")
            sec_shd.set(qn("w:val"), "clear")
            sec_shd.set(qn("w:color"), "auto")
            sec_shd.set(qn("w:fill"), "EFF6FF")
            sec_pPr.append(sec_shd)
            sec_para.paragraph_format.space_before = Pt(8)
            sec_para.paragraph_format.space_after = Pt(6)

        # Fields table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for field in fields:
            label = field.get("label", field.get("key", ""))
            value = field.get("value") or ""
            display = value if value.strip() else "________________"
            row = table.add_row()
            # Label cell
            label_cell = row.cells[0]
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = GRAY
            label_cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            # Value cell
            val_cell = row.cells[1]
            val_para = val_cell.paragraphs[0]
            val_run = val_para.add_run(display)
            val_run.font.size = Pt(11)
            val_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        # Remove table borders, keep bottom line per row
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBdr = OxmlElement("w:tcBdr")
                for side in ["top", "left", "right"]:
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:val"), "nil")
                    tcBdr.append(el)
                bot = OxmlElement("w:bottom")
                bot.set(qn("w:val"), "single")
                bot.set(qn("w:sz"), "4")
                bot.set(qn("w:color"), "E2E8F0")
                tcBdr.append(bot)
                tcPr.append(tcBdr)

        doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(f"Влада на Република Северна Македонија  •  е-Влада Портал  •  {today}")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = GRAY

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def service_exists_by_service_id(db: Session, service_id: int) -> bool:
    return db.query(Service.id).filter(Service.id == service_id).first() is not None


def get_template_by_service_id(db: Session, service_id: int) -> ServiceDocumentTemplate | None:
    return (
        db.query(ServiceDocumentTemplate)
        .filter(ServiceDocumentTemplate.service_id == service_id)
        .first()
    )


def upsert_template(db: Session, payload: ServiceDocumentTemplateCreate) -> ServiceDocumentTemplate:
    template = get_template_by_service_id(db, payload.service_id)
    if template is None:
        template = ServiceDocumentTemplate(**payload.model_dump())
        db.add(template)
    else:
        data = payload.model_dump()
        for key, value in data.items():
            setattr(template, key, value)

    db.commit()
    db.refresh(template)
    return template


def apply_template_update(
    template: ServiceDocumentTemplate,
    payload: ServiceDocumentTemplateUpdate,
) -> ServiceDocumentTemplate:
    data = payload.model_dump(exclude_unset=True)
    for key, value in data.items():
        setattr(template, key, value)
    return template


def create_blank_template_for_service(db: Session, service: Service) -> ServiceDocumentTemplate:
    existing_template = get_template_by_service_id(db, service.id)
    if existing_template is not None:
        return existing_template

    template = ServiceDocumentTemplate(
        service_id=service.id,
        title=f"{service.name} - формулар за апликација",
        template_type="json",
        template_body=build_default_template_body(service.name),
        is_active=True,
    )
    db.add(template)
    db.commit()
    db.refresh(template)
    return template


def get_template_by_service_name(db: Session, service_name: str) -> ServiceDocumentTemplate | None:
    return (
        db.query(ServiceDocumentTemplate)
        .join(Service, ServiceDocumentTemplate.service_id == Service.id)
        .filter(Service.name == service_name)
        .first()
    )


def detect_template_from_uploaded_document(
    db: Session,
    filename: str,
    document_text: str,
) -> ServiceDocumentTemplate | None:
    filename_match = re.match(r"^(?P<service_id>\d+?)-application-form(?:-filled)?\.txt$", filename)
    if filename_match:
        template = get_template_by_service_id(db, int(filename_match.group("service_id")))
        if template is not None:
            return template

    first_non_empty_line = next((line.strip() for line in document_text.splitlines() if line.strip()), "")
    if first_non_empty_line:
        template = get_template_by_service_name(db, first_non_empty_line)
        if template is not None:
            return template

    return None